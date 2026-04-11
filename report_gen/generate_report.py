"""
CvJachai Project Scan Report Generator
Run from the project root:
    python report_gen/generate_report.py
Output: report_gen/CvJachai_Scan_Report.docx
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime
import os

OUTPUT_PATH = os.path.join(os.path.dirname(__file__), "CvJachai_Scan_Report.docx")

# ── Colour palette ──────────────────────────────────────────────────────────
C_DARK    = RGBColor(0x1E, 0x29, 0x3B)   # header bg
C_ACCENT  = RGBColor(0x38, 0x82, 0xF4)   # blue accent
C_GREEN   = RGBColor(0x22, 0xC5, 0x5E)
C_YELLOW  = RGBColor(0xEA, 0xB3, 0x08)
C_RED     = RGBColor(0xEF, 0x44, 0x44)
C_WHITE   = RGBColor(0xFF, 0xFF, 0xFF)
C_LIGHT   = RGBColor(0xF1, 0xF5, 0xF9)
C_TEXT    = RGBColor(0x1E, 0x29, 0x3B)
C_MUTED   = RGBColor(0x64, 0x74, 0x8B)

HEX_DARK   = "1E293B"
HEX_ACCENT = "3882F4"
HEX_LIGHT  = "F1F5F9"
HEX_WHITE  = "FFFFFF"
HEX_GREEN  = "22C55E"
HEX_YELLOW = "EAB308"
HEX_RED    = "EF4444"
HEX_ORANGE = "F97316"
HEX_SLATE  = "E2E8F0"


# ── XML helpers ─────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color: str):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)


def set_cell_borders(cell, top=None, bottom=None, left=None, right=None):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side, val in [("top", top), ("bottom", bottom), ("left", left), ("right", right)]:
        if val:
            el = OxmlElement(f"w:{side}")
            el.set(qn("w:val"),   val.get("val", "single"))
            el.set(qn("w:sz"),    val.get("sz",  "4"))
            el.set(qn("w:color"), val.get("color", "auto"))
            tcBorders.append(el)
    tcPr.append(tcBorders)


def set_row_height(row, twips: int):
    tr   = row._tr
    trPr = tr.get_or_add_trPr()
    trH  = OxmlElement("w:trHeight")
    trH.set(qn("w:val"),  str(twips))
    trH.set(qn("w:hRule"), "atLeast")
    trPr.append(trH)


def no_space_para(para):
    pPr = para._p.get_or_add_pPr()
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:before"), "0")
    spacing.set(qn("w:after"),  "0")
    pPr.append(spacing)


# ── Style helpers ────────────────────────────────────────────────────────────

def add_heading(doc, text, level=1, color=C_DARK, space_before=18, space_after=6):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(space_before)
    para.paragraph_format.space_after  = Pt(space_after)
    run  = para.add_run(text)
    run.bold      = True
    run.font.color.rgb = color
    run.font.size = Pt({1: 18, 2: 14, 3: 12, 4: 11}.get(level, 11))
    return para


def add_body(doc, text, italic=False, color=C_TEXT, size=10, space_after=4):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after  = Pt(space_after)
    run  = para.add_run(text)
    run.italic         = italic
    run.font.size      = Pt(size)
    run.font.color.rgb = color
    return para


def add_bullet(doc, text, color=C_TEXT, size=10):
    para = doc.add_paragraph(style="List Bullet")
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after  = Pt(2)
    run  = para.add_run(text)
    run.font.size      = Pt(size)
    run.font.color.rgb = color
    return para


def add_divider(doc, color_hex=HEX_SLATE):
    para = doc.add_paragraph()
    no_space_para(para)
    pPr  = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot  = OxmlElement("w:bottom")
    bot.set(qn("w:val"),   "single")
    bot.set(qn("w:sz"),    "4")
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), color_hex)
    pBdr.append(bot)
    pPr.append(pBdr)
    para.paragraph_format.space_after = Pt(6)


def add_colored_table(doc, headers, rows, header_bg=HEX_DARK, row_alt_bg=HEX_LIGHT,
                       col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"

    # Header row
    hrow = table.rows[0]
    set_row_height(hrow, 400)
    for i, h in enumerate(headers):
        cell = hrow.cells[i]
        set_cell_bg(cell, header_bg)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p    = cell.paragraphs[0]
        no_space_para(p)
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after  = Pt(4)
        run  = p.add_run(h)
        run.bold           = True
        run.font.color.rgb = C_WHITE
        run.font.size      = Pt(9)

    # Data rows
    for ri, row_data in enumerate(rows):
        row   = table.rows[ri + 1]
        bg    = HEX_WHITE if ri % 2 == 0 else row_alt_bg
        set_row_height(row, 340)
        for ci, val in enumerate(row_data):
            cell = row.cells[ci]
            set_cell_bg(cell, bg)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p    = cell.paragraphs[0]
            no_space_para(p)
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after  = Pt(3)
            run  = p.add_run(str(val))
            run.font.size      = Pt(9)
            run.font.color.rgb = C_TEXT

    # Column widths
    if col_widths:
        for ci, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[ci].width = Inches(w)

    doc.add_paragraph()  # spacing after table
    return table


def add_badge_row(doc, items):
    """Add a row of coloured badge bullets."""
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(4)
    para.paragraph_format.space_after  = Pt(8)
    for label, hex_col in items:
        run = para.add_run(f"  {label}  ")
        run.font.size      = Pt(9)
        run.font.color.rgb = C_WHITE
        run.font.bold      = True
        # Fake highlight via shading isn't native in runs; add with space separator
        # Use Unicode block to simulate badge
        sep = para.add_run("   ")
        sep.font.size = Pt(9)
    return para


# ── Cover page ───────────────────────────────────────────────────────────────

def build_cover(doc):
    # Big header block
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell  = table.rows[0].cells[0]
    set_cell_bg(cell, HEX_DARK)
    set_row_height(table.rows[0], 2200)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    no_space_para(p)

    r = p.add_run("CvJachai")
    r.bold = True; r.font.size = Pt(36); r.font.color.rgb = C_WHITE

    p2 = cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    no_space_para(p2)
    r2 = p2.add_run("Resume Intelligence Engine")
    r2.font.size = Pt(16); r2.font.color.rgb = C_ACCENT; r2.bold = True

    p3 = cell.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    no_space_para(p3)
    r3 = p3.add_run("Full Project Scan Report")
    r3.font.size = Pt(12); r3.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)

    doc.add_paragraph()

    # Meta info table
    meta = [
        ("Project",    "CvJachai"),
        ("Corpus",     "nobiulhaque/CvJachai"),
        ("Scan Date",  datetime.datetime.now().strftime("%d %B %Y")),
        ("Version",    "3.2.0"),
        ("Deployment", "Hugging Face Spaces (Docker)"),
    ]
    mt = doc.add_table(rows=len(meta), cols=2)
    mt.style = "Table Grid"
    for ri, (k, v) in enumerate(meta):
        bg = HEX_LIGHT if ri % 2 == 0 else HEX_WHITE
        kc = mt.rows[ri].cells[0]; vc = mt.rows[ri].cells[1]
        set_cell_bg(kc, HEX_DARK); set_cell_bg(vc, bg)
        kc.width = Inches(1.8); vc.width = Inches(4.2)
        kp = kc.paragraphs[0]; no_space_para(kp)
        kp.paragraph_format.space_before = Pt(3)
        kp.paragraph_format.space_after  = Pt(3)
        kr = kp.add_run(k)
        kr.bold = True; kr.font.size = Pt(9); kr.font.color.rgb = C_ACCENT
        vp = vc.paragraphs[0]; no_space_para(vp)
        vp.paragraph_format.space_before = Pt(3)
        vp.paragraph_format.space_after  = Pt(3)
        vr = vp.add_run(v)
        vr.font.size = Pt(9); vr.font.color.rgb = C_TEXT

    doc.add_page_break()


# ── Section 1 — Directory Structure ─────────────────────────────────────────

def build_structure(doc):
    add_heading(doc, "1.  Project Directory Structure", level=1, color=C_ACCENT)
    add_divider(doc)

    tree = [
        ("cvjachai/",                          "Project root"),
        ("  ├── .env",                         "Secrets: SECRET_KEY, API_KEY, DB_PATH"),
        ("  ├── Dockerfile",                   "Production container → HF Spaces (port 7860)"),
        ("  ├── manage.py",                    "Django management entry-point"),
        ("  ├── requirements.txt",             "18 Python dependencies"),
        ("  ├── db.sqlite3",                   "Local dev DB (not used in prod)"),
        ("  ├── core/",                        "Django project config package"),
        ("  │   ├── settings.py",              "All Django + DRF + JWT configuration"),
        ("  │   ├── urls.py",                  "Root URL routing (7 URL patterns)"),
        ("  │   └── wsgi.py",                  "WSGI application entry-point"),
        ("  ├── api/",                         "Main Django app"),
        ("  │   ├── models.py",                "Profile model (OneToOne → User)"),
        ("  │   ├── serializers.py",           "ResumeUploadSerializer with full validation"),
        ("  │   ├── migrations/0001_initial.py","Initial DB schema migration"),
        ("  │   └── views/",                   "All API view modules"),
        ("  │       ├── auth_views.py",        "Signup / Login (custom email-JWT)"),
        ("  │       ├── classify_view.py",     "Core 3-stage resume ranking pipeline"),
        ("  │       ├── optimize_view.py",     "ATS resume optimizer (Groq 70B)"),
        ("  │       ├── profile_view.py",      "GET / PATCH user profile"),
        ("  │       └── info_view.py",         "Public API documentation endpoint"),
        ("  ├── engine/",                      "AI inference & text processing"),
        ("  │   ├── model.py",                 "LightGBM/CatBoost model loader & predictor"),
        ("  │   ├── pretrained_classifier.py", "Lazy HuggingFace NLI zero-shot reranker"),
        ("  │   ├── utils.py",                 "Text extraction: PDF, DOCX, TXT, OCR, ZIP"),
        ("  │   └── groq/",                    "Groq cloud AI sub-package"),
        ("  │       ├── client.py",            "GroqClient singleton (auto model discovery)"),
        ("  │       ├── ranker.py",            "GroqRanker — batch ranking (Llama 8B)"),
        ("  │       └── optimizer.py",         "GroqOptimizer — ATS rewrite (Llama 70B)"),
        ("  ├── models/",                      "Serialized ML artifacts"),
        ("  │   ├── best_classifier.pkl",      "Trained classifier (~2.1 MB)"),
        ("  │   ├── tfidf_vectorizer.pkl",     "TF-IDF vectorizer (~10 KB)"),
        ("  │   ├── scaler.pkl",               "Feature scaler (~29 KB)"),
        ("  │   └── model_metadata.json",      "Categories, skill list, accuracy (~15 KB)"),
        ("  └── data/",                        "(empty — training data not committed)"),
    ]

    t = doc.add_table(rows=len(tree), cols=2)
    t.style = "Table Grid"
    for ri, (path, desc) in enumerate(tree):
        bg  = HEX_LIGHT if ri % 2 == 0 else HEX_WHITE
        pc  = t.rows[ri].cells[0]
        dc  = t.rows[ri].cells[1]
        set_cell_bg(pc, bg); set_cell_bg(dc, bg)
        pc.width = Inches(2.8); dc.width = Inches(3.8)

        pp = pc.paragraphs[0]; no_space_para(pp)
        pp.paragraph_format.space_before = Pt(2)
        pp.paragraph_format.space_after  = Pt(2)
        pr = pp.add_run(path)
        pr.font.name = "Courier New"
        pr.font.size = Pt(8)
        pr.font.color.rgb = C_DARK
        if path.endswith("/"):
            pr.bold = True

        dp = dc.paragraphs[0]; no_space_para(dp)
        dp.paragraph_format.space_before = Pt(2)
        dp.paragraph_format.space_after  = Pt(2)
        dr = dp.add_run(desc)
        dr.font.size      = Pt(8)
        dr.font.color.rgb = C_MUTED

    doc.add_paragraph()


# ── Section 2 — Tech Stack ───────────────────────────────────────────────────

def build_tech_stack(doc):
    add_heading(doc, "2.  Technology Stack", level=1, color=C_ACCENT)
    add_divider(doc)

    rows = [
        ("Web Framework",      "Django 4.2 + DRF 3.14"),
        ("Authentication",     "JWT — djangorestframework-simplejwt 5.3.1"),
        ("Local ML Model",     "LightGBM / CatBoost (joblib-serialized)"),
        ("Feature Extraction", "TF-IDF (sklearn) + binary skill features + text stats"),
        ("Semantic Reranking", "HuggingFace cross-encoder/nli-MiniLM2-L6-H768 (zero-shot)"),
        ("Cloud AI (Rank)",    "Groq API — Llama 3.1 8B Instant"),
        ("Cloud AI (Opt.)",    "Groq API — Llama 3.3 70B Versatile"),
        ("OCR",                "RapidOCR (ONNX Runtime, CPU-only)"),
        ("PDF Extraction",     "PyPDF2 3.0.1"),
        ("DOCX Extraction",    "python-docx 1.1.0"),
        ("HTTP Server",        "Gunicorn 21.2.0 (--timeout 120)"),
        ("Database",           "SQLite3 (/tmp/db.sqlite3 in prod)"),
        ("Container",          "Docker (python:3.10-slim) → Hugging Face Spaces"),
    ]

    add_colored_table(doc,
        headers=["Layer", "Technology"],
        rows=rows,
        col_widths=[2.0, 4.6])


# ── Section 3 — API Endpoints ────────────────────────────────────────────────

def build_endpoints(doc):
    add_heading(doc, "3.  API Endpoints", level=1, color=C_ACCENT)
    add_divider(doc)

    rows = [
        ("1", "GET",        "/",                  "❌  Public", "Root — returns API info / live docs"),
        ("2", "GET",        "/api/",              "❌  Public", "Identical to root (duplicate route)"),
        ("3", "POST",       "/api/auth/signup",   "❌  Public", "Register new user (email + password)"),
        ("4", "POST",       "/api/auth/signin",   "❌  Public", "Login → returns access + refresh JWT + profile"),
        ("5", "GET/PATCH",  "/api/auth/profile",  "✅  JWT",    "View or update the authenticated user's profile"),
        ("6", "POST",       "/api/token/refresh/","❌  Public", "Exchange refresh token for new access token"),
        ("7", "POST",       "/api/classify",      "✅  JWT",    "Batch resume ranking (up to 1000 files)"),
        ("8", "POST",       "/api/optimize",      "✅  JWT",    "ATS resume rewrite powered by Llama 70B"),
    ]

    add_colored_table(doc,
        headers=["#", "Method", "URL", "Auth", "Description"],
        rows=rows,
        col_widths=[0.3, 0.8, 1.9, 0.9, 2.7])

    add_body(doc,
        "Note: /api/auth/profile counts as one URL pattern but serves two HTTP methods (GET & PATCH).",
        italic=True, color=C_MUTED, size=9)


# ── Section 4 — Classification Pipeline ─────────────────────────────────────

def build_pipeline(doc):
    add_heading(doc, "4.  Resume Classification Pipeline  (/api/classify)", level=1, color=C_ACCENT)
    add_divider(doc)

    stages = [
        ("Stage 1", "File Processing",
         "Uploaded files (PDF, DOCX, TXT, ZIP, PNG/JPG) are saved to a secure /tmp directory. "
         "ZIP archives are extracted and their contents iterated to find valid resume files. "
         "Unsupported formats raise a 400 Bad Request."),

        ("Stage 2", "Text Extraction (Multi-Threaded)",
         "Text is extracted concurrently using ThreadPoolExecutor (workers = CPU count). "
         "PDF → PyPDF2  |  DOCX → python-docx  |  TXT → UTF-8 read  |  Image → RapidOCR (ONNX). "
         "Files that yield no text are silently skipped and counted as failed_extractions."),

        ("Stage 3", "Local Keyword & Skill Analysis",
         "Every resume receives an initial_score: calculate_job_relevance() computes "
         "keyword-overlap between resume and job description (stopwords removed, 3+ char tokens only). "
         "calculate_skill_bonus() matches required skills and detects experience years. "
         "initial_score = relevance × 0.6 + skill_bonus × 0.4"),

        ("Stage 4", "Groq Cloud Ranking  (top 50)",
         "The 50 highest-scoring resumes are sent in chunks of 5 to Groq's Llama 3.1 8B model "
         "with a 2-second delay between chunks to respect rate limits. "
         "The model returns a JSON dict {filename: score 0.0–1.0}. "
         "If Groq is unavailable (no API_KEY), this stage is skipped entirely."),

        ("Stage 5", "Local NLI Semantic Reranking  (fallback, top 20)",
         "For resumes not scored by Groq (or if Groq fails), the top 20 candidates are passed "
         "through a HuggingFace NLI cross-encoder (cross-encoder/nli-MiniLM2-L6-H768). "
         "The model performs zero-shot entailment: 'This candidate is a perfect match for this job'. "
         "Remaining lower-ranked resumes keep their initial_score (PyTorch skipped for speed)."),

        ("Final", "Score Fusion & Result Assembly",
         "final_score = semantic × 0.50 + keyword_relevance × 0.30 + skill_bonus × 0.20\n"
         "Results are filtered by min_score (default 0.0) then sliced to top_k (default 5). "
         "Each result includes filename, final_score, semantic_score, keyword_relevance, "
         "skill_bonus, and the analysis_engine that judged it."),
    ]

    for badge, title, desc in stages:
        t = doc.add_table(rows=1, cols=2)
        t.style = "Table Grid"
        bc = t.rows[0].cells[0]; dc = t.rows[0].cells[1]
        set_cell_bg(bc, HEX_ACCENT); set_cell_bg(dc, HEX_LIGHT)
        bc.width = Inches(1.1); dc.width = Inches(5.5)
        bc.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        bp = bc.paragraphs[0]; no_space_para(bp)
        bp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        bp.paragraph_format.space_before = Pt(4)
        bp.paragraph_format.space_after  = Pt(4)
        br = bp.add_run(badge)
        br.bold = True; br.font.size = Pt(9); br.font.color.rgb = C_WHITE

        dp = dc.paragraphs[0]; no_space_para(dp)
        dp.paragraph_format.space_before = Pt(4)
        dp.paragraph_format.space_after  = Pt(4)
        dr = dp.add_run(f"{title}:  ")
        dr.bold = True; dr.font.size = Pt(9); dr.font.color.rgb = C_DARK
        dr2 = dp.add_run(desc)
        dr2.font.size = Pt(8.5); dr2.font.color.rgb = C_MUTED

        doc.add_paragraph().paragraph_format.space_after = Pt(3)


# ── Section 5 — Data Model ───────────────────────────────────────────────────

def build_data_model(doc):
    add_heading(doc, "5.  Data Model", level=1, color=C_ACCENT)
    add_divider(doc)

    add_heading(doc, "User  (Django built-in)", level=2, color=C_DARK, space_before=6)
    add_colored_table(doc,
        headers=["Field", "Type", "Note"],
        rows=[
            ("username",   "CharField", "Set to email value (email-as-username pattern)"),
            ("email",      "EmailField","Primary identifier / login credential"),
            ("password",   "CharField", "Hashed — min 8 chars, Django validators applied"),
            ("first_name", "CharField", "Optional"),
            ("last_name",  "CharField", "Optional"),
        ],
        col_widths=[1.4, 1.2, 4.0])

    add_heading(doc, "Profile  (OneToOne → User)", level=2, color=C_DARK, space_before=6)
    add_colored_table(doc,
        headers=["Field", "Type", "Note"],
        rows=[
            ("location",     "CharField", "max_length=255, blank/null allowed"),
            ("profession",   "CharField", "max_length=255, blank/null allowed"),
            ("skills",       "TextField", "Comma-separated string or free text"),
            ("company_name", "CharField", "max_length=255, blank/null allowed"),
            ("bio",          "TextField", "Free text biography"),
            ("phone_number", "CharField", "max_length=20, blank/null allowed"),
            ("website",      "URLField",  "blank/null allowed"),
        ],
        col_widths=[1.4, 1.2, 4.0])


# ── Section 6 — Configuration ────────────────────────────────────────────────

def build_config(doc):
    add_heading(doc, "6.  Configuration  (core/settings.py)", level=1, color=C_ACCENT)
    add_divider(doc)

    add_colored_table(doc,
        headers=["Setting", "Value", "Source"],
        rows=[
            ("SECRET_KEY",                  "From env",            ".env / environment variable"),
            ("DEBUG",                       "False (default)",     "env DEBUG — toggled per env"),
            ("ALLOWED_HOSTS",               "['*']",               "⚠ Wildcard — see Issues section"),
            ("DATABASE ENGINE",             "SQLite3",             ""),
            ("DATABASE NAME",               "/tmp/db.sqlite3",     "env DB_PATH overrides"),
            ("ACCESS_TOKEN_LIFETIME",       "2 hours",             "SIMPLE_JWT setting"),
            ("REFRESH_TOKEN_LIFETIME",      "1 day",               "SIMPLE_JWT setting"),
            ("ROTATE_REFRESH_TOKENS",       "False",               ""),
            ("FILE_UPLOAD_MAX_MEMORY_SIZE", "50 MB",               ""),
            ("DATA_UPLOAD_MAX_NUMBER_FILES","1000",                 "Supports large batch uploads"),
            ("Anon Throttle Rate",          "20 req/hour",         "DRF throttling"),
            ("User Throttle Rate",          "100 req/hour",        "DRF throttling"),
        ],
        col_widths=[2.4, 1.8, 2.4])


# ── Section 7 — Issues ───────────────────────────────────────────────────────

def build_issues(doc):
    add_heading(doc, "7.  Issues & Observations", level=1, color=C_ACCENT)
    add_divider(doc)

    issues = [
        ("🔴", "CRITICAL", HEX_RED,
         "ALLOWED_HOSTS = ['*']",
         "Wildcard host is a security risk in production. "
         "Should be restricted to the actual deployment domain e.g. your-space.hf.space."),

        ("🔴", "CRITICAL", HEX_RED,
         "django.contrib.staticfiles not in INSTALLED_APPS",
         "Django Admin panel will load without CSS styling. "
         "Low impact for a pure API but should be addressed if admin is used."),

        ("🟡", "MEDIUM", HEX_YELLOW,
         "Classifier loaded at module import time (startup)",
         "classify_view.py line 26 calls create_classifier() at module level. "
         "If the models/ directory is missing, the entire server fails to start instead of "
         "returning a graceful 503. Recommend lazy-loading with try/except."),

        ("🟡", "MEDIUM", HEX_YELLOW,
         "Model type label mismatch in get_model_info()",
         "engine/model.py returns 'model_type': 'CatBoost Classifier' but the module "
         "docstring and comments say LightGBM. The actual type should be confirmed from "
         "model_metadata.json and the label corrected."),

        ("🟡", "MEDIUM", HEX_YELLOW,
         "Blocking sleep in Groq ranker for large batches",
         "groq/ranker.py uses time.sleep(2.0) between every chunk of 5 resumes. "
         "Ranking 50 resumes = ~20 seconds of blocking sleep in a sync Django view. "
         "For large batches this can cause response timeouts. "
         "Consider async tasks (Celery) or reducing chunk delay."),

        ("🟡", "MEDIUM", HEX_YELLOW,
         "Hardcoded model name in optimize_view.py response",
         "Returns 'optimization_engine': 'Groq Llama 3.1 70B' as a static string, "
         "but the actual model is dynamically discovered via groq_base.optimizer_model. "
         "Should use the variable instead to stay accurate after auto-discovery."),

        ("🟡", "MEDIUM", HEX_YELLOW,
         "Pillow==9.5.0 has known CVEs",
         "requirements.txt pins Pillow at 9.5.0 which has security vulnerabilities. "
         "Recommend upgrading to Pillow>=10.3.0."),

        ("🟢", "LOW", HEX_GREEN,
         "Dead code block in serializers.py (lines 96–98)",
         "Two blank lines and an empty comment block sit between the extension check "
         "and the 'if errors:' block. Looks like removed size-validation logic was not "
         "cleaned up. Minor readability issue."),

        ("🟢", "LOW", HEX_GREEN,
         "output_gem.txt in project root",
         "A debug/scratch file is committed at the project root. "
         "Should be added to .gitignore and deleted from the repo."),

        ("🟢", "LOW", HEX_GREEN,
         ".github/ directory is empty — no CI/CD",
         "No GitHub Actions workflows exist. Adding a lint/test workflow "
         "would improve code quality and catch regressions automatically."),

        ("🟢", "LOW", HEX_GREEN,
         "PyPDF2 is deprecated upstream",
         "PyPDF2==3.0.1 is no longer maintained. The maintained fork is 'pypdf'. "
         "Not urgent, but worth migrating to avoid future security gaps."),

        ("🟢", "INFO", "3B82F6",
         "First classify request is slow (NLI model cold-start)",
         "SemanticReranker lazy-loads the HuggingFace model on first use. "
         "This is the correct design to avoid blocking startup, but the first request "
         "will take an additional ~5–10 seconds. Users should be informed via docs."),
    ]

    add_colored_table(doc,
        headers=["Sev.", "ID", "Issue", "Detail"],
        rows=[(icon, sev, title, detail) for icon, sev, _, title, detail in issues],
        col_widths=[0.35, 0.65, 2.2, 3.4])


# ── Section 8 — Dependencies ─────────────────────────────────────────────────

def build_dependencies(doc):
    add_heading(doc, "8.  Dependency Analysis  (requirements.txt)", level=1, color=C_ACCENT)
    add_divider(doc)

    rows = [
        ("django==4.2",                    "Core",    "Django web framework — stable LTS release"),
        ("djangorestframework==3.14.0",    "Core",    "DRF — REST API toolkit"),
        ("djangorestframework-simplejwt",  "Auth",    "JWT authentication"),
        ("python-dotenv==1.0.0",           "Config",  "Loads .env file into environment"),
        ("groq==1.1.0",                    "AI",      "Groq cloud LLM client"),
        ("gunicorn==21.2.0",               "Server",  "Production WSGI HTTP server"),
        ("torch==2.0.1",                   "AI ⚠",   "Heavy (~800 MB). Used only by transformers NLI fallback."),
        ("transformers==4.35.2",           "AI",      "HuggingFace NLI zero-shot reranker model"),
        ("rapidocr-onnxruntime==1.3.8",    "OCR",     "CPU-based OCR — good fit for HF Spaces (no GPU)"),
        ("python-docx==1.1.0",             "Extract", "DOCX text extraction"),
        ("pdfminer.six==20221105",         "Extract", "Installed but PyPDF2 is used for extraction"),
        ("Pillow==9.5.0",                  "Image ⚠", "Image processing — has known CVEs, upgrade to 10.x"),
        ("numpy<2.0.0",                    "ML",      "Pinned below 2.0 — correct for current ML ecosystem"),
        ("joblib==1.3.2",                  "ML",      "Model serialization / deserialization"),
        ("PyPDF2==3.0.1",                  "Extract ⚠","Deprecated; 'pypdf' is the maintained fork"),
        ("scikit-learn==1.3.0",            "ML",      "TF-IDF vectorizer + scaler"),
        ("catboost==1.2.2",                "ML ⚠",   "Both CatBoost & LightGBM installed; only one used"),
        ("lightgbm==4.1.0",               "ML ⚠",   "Both CatBoost & LightGBM installed; only one used"),
    ]

    add_colored_table(doc,
        headers=["Package", "Category", "Notes"],
        rows=rows,
        col_widths=[2.2, 0.9, 3.5])

    add_body(doc,
        "⚠  pdfminer.six appears in requirements.txt but PyPDF2 is the active extraction library. "
        "pdfminer may be unused dead weight adding to Docker image size.",
        italic=True, color=C_MUTED, size=9)


# ── Section 9 — Summary scorecard ────────────────────────────────────────────

def build_scorecard(doc):
    add_heading(doc, "9.  Health Scorecard", level=1, color=C_ACCENT)
    add_divider(doc)

    rows = [
        ("Code Organisation",   "✅  Excellent",  "Clear separation: api/, engine/, models/, core/"),
        ("Authentication",      "✅  Good",        "JWT with email-as-username, profile auto-created"),
        ("Input Validation",    "✅  Good",        "Full serializer validation on classify endpoint"),
        ("Error Handling",      "✅  Good",        "Try/except across all stages with HTTP codes"),
        ("AI Ranking Quality",  "✅  Excellent",  "Groq → NLI → Keyword 3-tier fallback system"),
        ("Deployment",          "✅  Good",        "Docker, gunicorn, /tmp DB, HF Spaces ready"),
        ("Security",            "⚠️  Fair",        "ALLOWED_HOSTS wildcard; Pillow CVE"),
        ("Performance",         "⚠️  Fair",        "Blocking sleep in ranker; cold-start NLI load"),
        ("Dependencies",        "⚠️  Fair",        "Duplicate ML libs; deprecated PyPDF2; unused pdfminer"),
        ("CI/CD",               "❌  Missing",     "No GitHub Actions workflows in .github/"),
        ("Test Coverage",       "❌  Missing",     "No test files found in project"),
    ]

    add_colored_table(doc,
        headers=["Area", "Rating", "Comment"],
        rows=rows,
        col_widths=[1.8, 1.2, 3.6])


# ── Main ─────────────────────────────────────────────────────────────────────

def main():
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin    = Cm(1.8)
        section.bottom_margin = Cm(1.8)
        section.left_margin   = Cm(2.0)
        section.right_margin  = Cm(2.0)

    # Default font
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(10)

    build_cover(doc)
    build_structure(doc)
    doc.add_page_break()
    build_tech_stack(doc)
    build_endpoints(doc)
    doc.add_page_break()
    build_pipeline(doc)
    doc.add_page_break()
    build_data_model(doc)
    build_config(doc)
    doc.add_page_break()
    build_issues(doc)
    doc.add_page_break()
    build_dependencies(doc)
    build_scorecard(doc)

    doc.save(OUTPUT_PATH)
    print(f"\nDone! Report saved: {OUTPUT_PATH}\n")


if __name__ == "__main__":
    main()
